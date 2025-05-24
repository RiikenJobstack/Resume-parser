import os
import io
import base64
import re
from typing import List
from io import BytesIO

import pdfplumber
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes

import docx
from app.config import settings


class DocumentService:
    """Service for processing document files"""

    @staticmethod
    def extract_text_from_file(file_buffer: BytesIO, file_extension: str) -> str:
        ext = file_extension.lower()
        if ext == '.pdf':
            text = DocumentService.extract_text_from_pdf(file_buffer)
            if not text.strip():
                # Try OCR fallback if no text extracted
                text = DocumentService.extract_text_from_pdf_ocr(file_buffer)
            return text

        elif ext == '.docx':
            return DocumentService.extract_text_from_docx(file_buffer)

        elif ext == '.txt':
            return file_buffer.read().decode('utf-8', errors='replace')

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    @staticmethod
    def extract_text_from_pdf(buffer: BytesIO) -> str:
        """Try pdfplumber first, then fallback to PyMuPDF if needed"""
        text = ""
        try:
            buffer.seek(0)
            with pdfplumber.open(buffer) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if page_text:
                        text += page_text + "\n\n"

                    # Extract tables
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            text += "\n" + "\n".join([" | ".join([str(cell or "") for cell in row]) for row in table]) + "\n"
            if text.strip():
                return text

            # Fallback to PyMuPDF if no text extracted
            buffer.seek(0)
            return DocumentService.extract_text_from_pdf_pymupdf(buffer)

        except Exception as e:
            # Fallback to PyMuPDF if pdfplumber fails
            buffer.seek(0)
            return DocumentService.extract_text_from_pdf_pymupdf(buffer)

    @staticmethod
    def extract_text_from_pdf_pymupdf(buffer: BytesIO) -> str:
        """Extract text using PyMuPDF (fitz) as a fallback"""
        text = ""
        try:
            buffer.seek(0)
            pdf_doc = fitz.open(stream=buffer.read(), filetype="pdf")
            for page in pdf_doc:
                page_text = page.get_text("text")  # or "blocks", "html" if needed
                if page_text:
                    text += page_text + "\n\n"
            return text
        except Exception as e:
            return ""

    @staticmethod
    def extract_text_from_pdf_ocr(buffer: BytesIO) -> str:
        """
        Use OCR (Tesseract) to extract text from scanned PDFs or images embedded in PDFs
        Requires `pdf2image` and `pytesseract` installed and Tesseract-OCR setup on the system.
        """
        try:
            buffer.seek(0)
            images = convert_from_bytes(buffer.read())
            text = ""
            for img in images:
                text += pytesseract.image_to_string(img) + "\n\n"
            return text
        except Exception as e:
            return ""

    @staticmethod
    def extract_text_from_docx(buffer: BytesIO) -> str:
        """
        Extract text from DOCX using python-docx for more robust parsing
        """
        try:
            buffer.seek(0)
            doc = docx.Document(buffer)
            full_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    full_text.append(row_text)

            return "\n\n".join(full_text)

        except Exception as e:
            # fallback to docx2txt
            import docx2txt
            buffer.seek(0)
            return docx2txt.process(buffer)

    @staticmethod
    def preprocess_text(text: str) -> str:
        if not text:
            return ""

        # Replace multiple whitespace with single space, but keep line breaks for structure
        # First normalize line breaks
        text = re.sub(r'\r\n|\r', '\n', text)

        # Replace multiple blank lines with max two
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        # Normalize bullet points (•, -, *, etc.)
        text = re.sub(r'[\u2022•*]\s*', '- ', text)

        # Add space after colons in headers, if missing
        text = re.sub(r'([A-Z][A-Z\s]+):?(\S)', r'\1: \2', text)

        # Add extra newlines before all-uppercase headers (common in resumes)
        text = re.sub(r'\n([A-Z][A-Z\s]{2,})\n', r'\n\n\1\n\n', text)

        # Strip trailing and leading spaces on lines
        text = "\n".join(line.strip() for line in text.splitlines())

        return text.strip()

    @staticmethod
    def classify_resume_complexity(text: str) -> str:
        token_count = len(text.split())

        has_tables = '|' in text and text.count('|') > 10

        technical_keywords = [
            'algorithm', 'framework', 'architecture', 'infrastructure',
            'kubernetes', 'docker', 'aws', 'azure', 'cloud', 'devops',
            'backend', 'frontend', 'fullstack', 'machine learning', 'ai',
            'microservices', 'distributed', 'scalable', 'optimization'
        ]
        technical_count = sum(1 for word in technical_keywords if word.lower() in text.lower())

        complexity_score = (
            (token_count > 1000) * 1 +
            has_tables * 2 +
            (technical_count > 5) * 2
        )

        return 'complex' if complexity_score >= 3 else 'simple'
